import os
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def parse_max_size(text):
    parts = text.split(", Max:")
    if len(parts) > 1:
        max_size_part = parts[1].split(",")[0].strip()
        return int(max_size_part)
    return None

def get_file_path(directory, filename):
    for file in os.listdir(directory):
        if file == filename:
            return os.path.join(directory, file)
    return None

def parse_file_lines(file_path):
    best = -1
    generation = -1
    lc = 1
    with open(file_path, "r") as file:
        for line in file:
            line = line.strip()
            if line:
                cur_max = parse_max_size(line)
                if cur_max is not None and cur_max > best:
                    best = cur_max
                    generation = lc
            if "Gen:" in line:
                lc += 1
    return best, generation

def get_result(dataset_name):
    file_path = get_file_path("./dimacs_results/", f"{dataset_name}_result.txt")
    if file_path:
        return parse_file_lines(file_path)
    return [-1, -1]

def parse_line(parts, results):
    best, gen = results
    parts.append(best)
    parts.append(gen)
    bk = int(re.search(r"\d+", parts[3]).group())
    if bk == best:
        parts.append("E")
    elif bk > best:
        parts.append("L")
    else:
        parts.append("GG")
    return parts

def save_results_to_file(data, output_file, less_datasets, greater_datasets):
    with open(output_file, "w") as file:
        file.write("Dimacs Results\n")
        file.write("------------------\n")
        for line in data:
            file.write(" ".join(map(str, line)) + "\n")
        
        file.write("\n\nLess Values Found:\n")
        file.write("------------------\n")
        for i, dataset in enumerate(less_datasets, start=1):
            file.write(f"{i}. {dataset}\n")
        
        file.write("\n\nGreater Values Found:\n")
        file.write("----------------------\n")
        for i, dataset in enumerate(greater_datasets, start=1):
            file.write(f"{i}. {dataset}\n")

def generate_word_document(data, output_doc, less_datasets, greater_datasets):
    if os.path.exists(output_doc):
        os.remove(output_doc)

    doc = Document()
    doc.add_heading("Dimacs Results", level=1)
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, header in enumerate(["Dataset", "|V|", "|E|", "ω(G)", "Best", "Generation", "Status"]):
        header_row.cells[i].text = header
        header_row.cells[i].paragraphs[0].runs[0].font.bold = True

    for row_data in data:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data)
        
        status = row_data[-1]
        color = "FFFFFF"
        if status == "GG":
            color = "00FF00"
        elif status == "L":
            color = "FF0000"
        
        for cell in row:
            cell_element = cell._tc
            cell_properties = cell_element.get_or_add_tcPr()
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), color)
            cell_properties.append(shade)

    doc.add_heading("Less Values Found", level=1)
    for i, dataset in enumerate(less_datasets, start=1):
        doc.add_paragraph(f"{i}. {dataset}")

    doc.add_heading("Greater Values Found", level=1)
    for i, dataset in enumerate(greater_datasets, start=1):
        doc.add_paragraph(f"{i}. {dataset}")

    doc.save(output_doc)

def gen_raw_report(input_file, text_output, word_output):
    results = []
    with open(input_file, "r") as file:
        next(file)
        for line in file:
            parts = line.strip().split()
            if parts:
                results.append(parse_line(parts, get_result(parts[0])))

    os.makedirs(os.path.dirname(text_output), exist_ok=True)
    less_datasets = []
    greater_datasets = []
    for row in results:
        if row[-1] == "L":
            less_datasets.append(row[0])
        elif row[-1] == "GG":
            greater_datasets.append(row[0])

    save_results_to_file(results, text_output, less_datasets, greater_datasets)
    generate_word_document(results, word_output, less_datasets, greater_datasets)

    print("\nLess Values Found:")
    for i, dataset in enumerate(less_datasets, start=1):
        print(f"{i}. {dataset}")

    print("\nGreater Values Found:")
    for i, dataset in enumerate(greater_datasets, start=1):
        print(f"{i}. {dataset}")

if __name__ == "__main__":
    input_file = "./best_known/dimacs_best_known.txt"
    text_output = "./results/dimacs_results.txt"
    word_output = "./results/dimacs_results.docx"
    gen_raw_report(input_file, text_output, word_output)
