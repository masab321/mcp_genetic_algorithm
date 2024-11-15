import os
from docx import Document

def extract_clique_info(file_path):
    with open(file_path, 'r') as file:
        # Read the last three lines of the file
        lines = file.readlines()[-3:]
        
        # Initialize variables
        filename = None
        max_clique_size = None
        clique_vertices = None

        # Look for the lines with the filename, max clique size, and clique vertices
        for line in lines:
            if "Clique result for file:" in line:
                filename = line.split("Clique result for file:")[-1].strip()
                filename = filename.split("/")[-1]  # Extract just the filename
            elif "Max clique size:" in line:
                max_clique_size = int(line.split("Max clique size:")[-1].strip())
            elif "Vertices in max clique:" in line:
                clique_vertices = line.split("Vertices in max clique:")[-1].strip()
        
        return filename, max_clique_size, clique_vertices

def main():
    # Directory containing result files
    directory = "Dimacs_results"
    
    # Create a Word document
    doc = Document()
    doc.add_heading('Max Clique Results', level=1)
    
    # Add a table with two columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filename'
    hdr_cells[1].text = 'Max Clique Size'
    
    # Loop through each file in the directory and extract information
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        
        if os.path.isfile(file_path):
            clq_filename, max_clique_size, clique_vertices = extract_clique_info(file_path)
            
            if clq_filename and max_clique_size and clique_vertices:
                # Print filename and max clique size
                print(f"Filename: {clq_filename}, Max Clique Size: {max_clique_size}")
                
                # Uncomment the line below to print the clique vertices as well
                # print(f"Filename: {clq_filename}, Max Clique Size: {max_clique_size}, Clique: {clique_vertices}")
                
                # Add data to the table in the Word document
                row_cells = table.add_row().cells
                row_cells[0].text = clq_filename
                row_cells[1].text = str(max_clique_size)
    
    # Save the document
    doc.save('Max_Clique_Results.docx')
    print("Data saved to Max_Clique_Results.docx")

if __name__ == "__main__":
    main()
